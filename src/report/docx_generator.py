from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt


# ============================================================
# REPORT CONFIGURATION
# ============================================================

REPORT_SECTIONS = [
    "Reporting Period",
    "Narrative Summary",
    "Case Summary",
    "Reaction Analysis",
    "Serious Cases",
    "Trends",
    "History of Actions",
    "Case Index",
]


# ============================================================
# MARKDOWN HELPERS
# ============================================================

def clean_markdown_inline(text: str) -> str:
    """
    Remove common Markdown formatting while preserving text.
    """

    if not text:
        return ""

    # Bold
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)

    # Italic
    text = re.sub(r"\*(.*?)\*", r"\1", text)

    # Inline code
    text = re.sub(r"`(.*?)`", r"\1", text)

    # Markdown links:
    # [text](url) -> text
    text = re.sub(
        r"\[([^\]]+)\]\([^)]+\)",
        r"\1",
        text,
    )

    return text.strip()


def is_heading(line: str) -> bool:
    """
    Check whether a line is a Markdown heading.
    """

    return bool(
        re.match(
            r"^\s*#{1,6}\s+.+",
            line,
        )
    )


def heading_level(line: str) -> int:
    """
    Return Markdown heading level.
    """

    match = re.match(
        r"^\s*(#{1,6})\s+",
        line,
    )

    if not match:
        return 0

    return len(match.group(1))


def heading_text(line: str) -> str:
    """
    Extract heading text from a Markdown heading.
    """

    text = re.sub(
        r"^\s*#{1,6}\s+",
        "",
        line,
    )

    return clean_markdown_inline(text)


def is_bullet(line: str) -> bool:
    """
    Detect Markdown bullet items.
    """

    return bool(
        re.match(
            r"^\s*[-*+]\s+",
            line,
        )
    )


def bullet_text(line: str) -> str:
    """
    Extract bullet text.
    """

    text = re.sub(
        r"^\s*[-*+]\s+",
        "",
        line,
    )

    return clean_markdown_inline(text)


def is_numbered_list(line: str) -> bool:
    """
    Detect Markdown numbered lists.
    """

    return bool(
        re.match(
            r"^\s*\d+\.\s+",
            line,
        )
    )


def numbered_list_text(line: str) -> str:
    """
    Extract numbered-list text.
    """

    text = re.sub(
        r"^\s*\d+\.\s+",
        "",
        line,
    )

    return clean_markdown_inline(text)


# ============================================================
# DOCUMENT STYLE
# ============================================================

def configure_document(document: Document) -> None:
    """
    Configure global document styles.
    """

    styles = document.styles

    # --------------------------------------------------------
    # Normal text
    # --------------------------------------------------------

    normal_style = styles["Normal"]

    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15

    # --------------------------------------------------------
    # Title
    # --------------------------------------------------------

    title_style = styles["Title"]

    title_style.font.name = "Arial"
    title_style.font.size = Pt(20)
    title_style.font.bold = True

    # --------------------------------------------------------
    # Heading 1
    # --------------------------------------------------------

    heading1 = styles["Heading 1"]

    heading1.font.name = "Arial"
    heading1.font.size = Pt(15)
    heading1.font.bold = True

    heading1.paragraph_format.space_before = Pt(14)
    heading1.paragraph_format.space_after = Pt(7)

    # --------------------------------------------------------
    # Heading 2
    # --------------------------------------------------------

    heading2 = styles["Heading 2"]

    heading2.font.name = "Arial"
    heading2.font.size = Pt(13)
    heading2.font.bold = True

    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(5)

    # --------------------------------------------------------
    # Heading 3
    # --------------------------------------------------------

    heading3 = styles["Heading 3"]

    heading3.font.name = "Arial"
    heading3.font.size = Pt(11.5)
    heading3.font.bold = True


# ============================================================
# PAGE CONFIGURATION
# ============================================================

def configure_page(document: Document) -> None:
    """
    Configure page margins.
    """

    for section in document.sections:

        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


# ============================================================
# ADD TITLE
# ============================================================

def add_report_title(document: Document) -> None:
    """
    Add the main report title.
    """

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(
        "PADER Safety Analysis Report"
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)

    paragraph.space_after = Pt(4)

    subtitle = document.add_paragraph()

    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = subtitle.add_run(
        "AI-Assisted Deterministic Analysis and Evidence-Grounded Reporting"
    )

    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.italic = True

    subtitle.paragraph_format.space_after = Pt(14)


# ============================================================
# ADD PARAGRAPH
# ============================================================

def add_paragraph(
    document: Document,
    text: str,
) -> None:
    """
    Add a normal paragraph.
    """

    text = clean_markdown_inline(text)

    if not text:
        return

    paragraph = document.add_paragraph()

    paragraph.add_run(text)


# ============================================================
# ADD BULLET
# ============================================================

def add_bullet(
    document: Document,
    text: str,
) -> None:
    """
    Add a bullet paragraph.
    """

    text = clean_markdown_inline(text)

    if not text:
        return

    paragraph = document.add_paragraph(
        style="List Bullet"
    )

    paragraph.add_run(text)


# ============================================================
# ADD NUMBERED ITEM
# ============================================================

def add_numbered_item(
    document: Document,
    text: str,
) -> None:
    """
    Add a numbered list item.
    """

    text = clean_markdown_inline(text)

    if not text:
        return

    paragraph = document.add_paragraph(
        style="List Number"
    )

    paragraph.add_run(text)


# ============================================================
# ADD HORIZONTAL RULE
# ============================================================

def add_horizontal_rule(
    document: Document,
) -> None:
    """
    Add a simple horizontal separator.
    """

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)

    run = paragraph.add_run(
        "─" * 80
    )

    run.font.size = Pt(7)


# ============================================================
# PARSE MARKDOWN
# ============================================================

def add_markdown_to_document(
    document: Document,
    markdown_text: str,
) -> None:
    """
    Convert Markdown content into DOCX paragraphs.
    """

    lines = markdown_text.splitlines()

    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:

        nonlocal paragraph_buffer

        if not paragraph_buffer:
            return

        text = " ".join(
            item.strip()
            for item in paragraph_buffer
        ).strip()

        if text:
            add_paragraph(
                document,
                text,
            )

        paragraph_buffer = []

    for raw_line in lines:

        line = raw_line.rstrip()

        # ----------------------------------------------------
        # Empty line
        # ----------------------------------------------------

        if not line.strip():

            flush_paragraph()

            continue

        # ----------------------------------------------------
        # Markdown heading
        # ----------------------------------------------------

        if is_heading(line):

            flush_paragraph()

            level = heading_level(line)

            text = heading_text(line)

            if not text:
                continue

            # Avoid making the document title twice.
            if (
                level == 1
                and text.lower()
                in {
                    "pader safety analysis report",
                    "pader report",
                }
            ):
                continue

            document.add_heading(
                text,
                level=min(level, 3),
            )

            continue

        # ----------------------------------------------------
        # Bullet
        # ----------------------------------------------------

        if is_bullet(line):

            flush_paragraph()

            add_bullet(
                document,
                bullet_text(line),
            )

            continue

        # ----------------------------------------------------
        # Numbered list
        # ----------------------------------------------------

        if is_numbered_list(line):

            flush_paragraph()

            add_numbered_item(
                document,
                numbered_list_text(line),
            )

            continue

        # ----------------------------------------------------
        # Horizontal rule
        # ----------------------------------------------------

        if re.match(
            r"^\s*([-*_]){3,}\s*$",
            line,
        ):

            flush_paragraph()

            add_horizontal_rule(
                document
            )

            continue

        # ----------------------------------------------------
        # Normal paragraph
        # ----------------------------------------------------

        paragraph_buffer.append(
            line.strip()
        )

    flush_paragraph()


# ============================================================
# FOOTER
# ============================================================

def add_footer(document: Document) -> None:
    """
    Add a simple footer to every page.
    """

    for section in document.sections:

        footer = section.footer

        paragraph = footer.paragraphs[0]

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            "PADER AI Engine"
        )

        run.font.name = "Arial"
        run.font.size = Pt(8)


# ============================================================
# GENERATE DOCX
# ============================================================

def generate_docx(
    markdown_path: Path,
    output_path: Path,
) -> Path:
    """
    Convert a Markdown report into a DOCX document.

    Parameters
    ----------
    markdown_path:
        Path to final_report.md.

    output_path:
        Destination path for final_report.docx.

    Returns
    -------
    Path
        Generated DOCX path.
    """

    markdown_path = Path(
        markdown_path
    )

    output_path = Path(
        output_path
    )

    if not markdown_path.exists():

        raise FileNotFoundError(
            f"Markdown report not found: "
            f"{markdown_path}"
        )

    markdown_text = (
        markdown_path.read_text(
            encoding="utf-8"
        )
    )

    if not markdown_text.strip():

        raise ValueError(
            "Markdown report is empty."
        )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    configure_document(
        document
    )

    configure_page(
        document
    )

    add_report_title(
        document
    )

    add_markdown_to_document(
        document,
        markdown_text,
    )

    add_footer(
        document
    )

    document.save(
        output_path
    )

    return output_path


# ============================================================
# COMMAND LINE ENTRY POINT
# ============================================================

def main() -> None:

    project_root = (
        Path(__file__)
        .resolve()
        .parents[2]
    )

    markdown_path = (
        project_root
        / "outputs"
        / "generated_report"
        / "final_report.md"
    )

    output_path = (
        project_root
        / "outputs"
        / "generated_report"
        / "final_report.docx"
    )

    print(
        "\n"
        + "=" * 70
    )

    print(
        "PADER DOCX REPORT GENERATOR"
    )

    print(
        "=" * 70
    )

    print(
        f"Source : {markdown_path}"
    )

    print(
        f"Output : {output_path}"
    )

    print(
        "\nGenerating DOCX..."
    )

    generated_path = generate_docx(
        markdown_path=markdown_path,
        output_path=output_path,
    )

    print(
        "\nDOCX generation completed."
    )

    print(
        f"Saved: {generated_path}"
    )

    print(
        "=" * 70
    )


if __name__ == "__main__":
    main()